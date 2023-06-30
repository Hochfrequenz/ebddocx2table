"""
Contains high level functions to process .docx files
"""
import itertools
import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Dict, Generator, Iterable, List, Optional, Tuple, Union

import attrs
from docx import Document  # type:ignore[import]
from docx.oxml import CT_P, CT_Tbl  # type:ignore[import]
from docx.table import Table, _Cell  # type:ignore[import]
from docx.text.paragraph import Paragraph  # type:ignore[import]

_logger = logging.getLogger(__name__)


def get_document(docx_file_path: Path) -> Document:
    """
    opens and returns the document specified in the docx_file_path using python-docx
    """
    with open(docx_file_path, "rb") as docx_file:
        source_stream = BytesIO(docx_file.read())
        # Originally I tried the recipe from
        # https://python-docx.readthedocs.io/en/latest/user/documents.html#opening-a-file-like-document
        # but then switched from StringIO to BytesIO (without explicit 'utf-8') because of:
        # UnicodeDecodeError: 'charmap' codec can't decode byte 0x81 in position 605: character maps to <undefined>
    try:
        document = Document(source_stream)
        _logger.info("Successfully read the file '%s'", docx_file_path)
        return document
    finally:
        source_stream.close()


def _get_tables_and_paragaphs(document: Document) -> Generator[Union[Table, Paragraph], None, None]:
    """
    Yields tables and paragraphs from the given document in the order in which they occur in the document.
    This is helpful because document.tables and document.paragraphs are de-coupled and give you no information which
    paragraph follows which table.
    """
    parent_elements = document.element.body
    for item in parent_elements.iterchildren():
        if isinstance(item, CT_P):
            yield Paragraph(item, document)
        elif isinstance(item, CT_Tbl):
            yield Table(item, document)


_ebd_key_pattern = re.compile(r"^E_\d{4}$")
_ebd_key_with_heading_pattern = re.compile(r"^(?P<key>E_\d{4})_(?P<title>.*)\s*$")


class TableNotFoundError(Exception):
    """
    an error that is raised when a requested table was not found
    """

    def __init__(self, ebd_key: str):
        self.ebd_key = ebd_key


_ebd_cell_pattern = re.compile(r"^(?:ja|nein)\s*(?:Ende|\d+)$")
"""
any EBD table shall contain at least one cell that matches this pattern
"""


def _cell_is_probably_from_an_ebd_cell(cell: _Cell) -> bool:
    if "" in cell.text:
        return True
    if cell.text in {"ja", "nein"}:
        return True
    if "à" in cell.text:
        # the rightarrow in wrong encoding
        return True
    if _ebd_cell_pattern.match(cell.text):
        return True
    if cell.text.strip().startswith("Cluster:") or cell.text.startswith("Hinweis:"):
        return True
    return False


def _table_is_an_ebd_table(table: Table) -> bool:
    """
    Returns true iff the table "looks like" an EB-Table.
    This is to distinguish between tables that are inside the same subsection that describes an EBD but are not part
    of the decision tree at all (e.g. in E_0406 the tables about Artikel-IDs).
    """
    for row in table.rows:
        try:
            for cell in row.cells:
                if _cell_is_probably_from_an_ebd_cell(cell):
                    return True
        except IndexError:  # don't ask me why this happens; It's the internals of python-docx
            continue
    return False


def get_ebd_docx_tables(docx_file_path: Path, ebd_key: str) -> List[Table]:
    """
    Opens the file specified in docx_file_path and returns the tables that relate to the given ebd_key.
    There might be more than 1 docx table for 1 EBD table.
    This is because of inconsistencies and manual editing during creation of the documents by EDI@Energy.
    Raises an TableNotFoundError if the table was not found.
    """
    if _ebd_key_pattern.match(ebd_key) is None:
        raise ValueError(f"The ebd_key '{ebd_key}' does not match {_ebd_key_pattern.pattern}")
    document = get_document(docx_file_path)

    is_inside_subsection_of_requested_table: bool = False
    tables: List[Table] = []
    tables_and_paragraphs = _get_tables_and_paragaphs(document)
    for table_or_paragraph in tables_and_paragraphs:
        if isinstance(table_or_paragraph, Paragraph):
            paragraph: Paragraph = table_or_paragraph
            # Assumptions:
            # 1. before each EbdTable there is a paragraph whose text starts with the respective EBD key
            # 2. there are no duplicates
            is_inside_subsection_of_requested_table = paragraph.text.startswith(ebd_key) or (
                is_inside_subsection_of_requested_table
                and not (
                    paragraph.text.strip().startswith("Es it das EBD") and paragraph.text.strip().endswith("zu nutzen.")
                )
            )
        if (
            isinstance(table_or_paragraph, Table)
            and is_inside_subsection_of_requested_table
            and _table_is_an_ebd_table(table_or_paragraph)
        ):
            table: Table = table_or_paragraph
            tables.append(table)
            # Now we have to check if the EBD table spans multiple pages and _maybe_ we have to collect more tables.
            # The funny thing is: Sometimes the authors create multiple tables split over multiple lines which belong
            # together, sometimes they create 1 proper table that spans multiple pages.
            # The latter case (1 docx table spanning >1 pages) is transparent to the extraction logic; i.e. python-docx
            # treats a single table that spans multiple pages just the same as a table on only 1 page.
            for next_item in tables_and_paragraphs:  # start iterating from where the outer loop paused
                if isinstance(next_item, Table):
                    # this is the case that the authors created multiple single tables on single adjacent pages
                    # if table_is_an_ebd_table(table):
                    if _table_is_an_ebd_table(next_item):
                        tables.append(next_item)
                elif isinstance(next_item, Paragraph):
                    if next_item.text.startswith("S_") or next_item.text.startswith("E_"):
                        # this is the case that the authors created 1 table that spans multiple pages
                        # and we're done collecting tables for this EBD key
                        break
                    continue
                else:
                    break  # inner loop because if no other table will follow
                    # we're done collecting the tables for this EBD key
        if is_inside_subsection_of_requested_table and len(tables) > 0:  # this means: we found the table
            # break the outer loop, too; no need to iterate any further
            break
    if not any(tables):
        raise TableNotFoundError(ebd_key=ebd_key)
    return tables


# pylint:disable=too-few-public-methods
@attrs.define(kw_only=True, frozen=True)
class EbdChapterInformation:
    """
    Contains information about where an EBD is located within the document.
    If the heading is e.g. "5.2.1" we denote this as:
    * chapter 5
    * section 2
    * subsection 1
    """

    chapter: int = attrs.field(
        validator=attrs.validators.and_(attrs.validators.instance_of(int), attrs.validators.ge(1))
    )
    chapter_title: Optional[str] = attrs.field(validator=attrs.validators.optional(attrs.validators.instance_of(str)))
    section: int = attrs.field(
        validator=attrs.validators.and_(attrs.validators.instance_of(int), attrs.validators.ge(1))
    )

    section_title: Optional[str] = attrs.field(validator=attrs.validators.optional(attrs.validators.instance_of(str)))
    subsection: int = attrs.field(
        validator=attrs.validators.and_(attrs.validators.instance_of(int), attrs.validators.ge(1))
    )

    subsection_title: Optional[str] = attrs.field(
        validator=attrs.validators.optional(attrs.validators.instance_of(str))
    )


def _enrich_paragraphs_with_sections(
    paragraphs: Iterable[Paragraph],
) -> Generator[Tuple[Paragraph, EbdChapterInformation], None, None]:
    """
    Yield each paragraph + the "Kapitel" in which it is found.
    """
    chapter_counter = itertools.count(start=1)
    chapter = 1
    chapter_title: Optional[str] = None
    section_counter = itertools.count(start=1)
    section = 1
    section_title: Optional[str] = None
    subsection_counter = itertools.count(start=1)
    subsection = 1
    subsection_title: Optional[str] = None
    for paragraph in paragraphs:
        match paragraph.style.style_id:
            case "berschrift1":
                chapter = next(chapter_counter)
                chapter_title = paragraph.text.strip()
                section_counter = itertools.count(start=1)
                section_title = None
                subsection_counter = itertools.count(start=1)
                subsection_title = None
            case "berschrift2":
                section = next(section_counter)
                section_title = paragraph.text.strip()
                subsection_counter = itertools.count(start=1)
                subsection_title = None
            case "berschrift3":
                subsection = next(subsection_counter)
                subsection_title = paragraph.text.strip()
        location = EbdChapterInformation(
            chapter=chapter,
            section=section,
            subsection=subsection,
            chapter_title=chapter_title,
            section_title=section_title,
            subsection_title=subsection_title,
        )
        _logger.debug("Handling Paragraph %i.%i.%i", chapter, section, subsection)
        yield paragraph, location


def get_all_ebd_keys(docx_file_path: Path) -> Dict[str, Tuple[str, EbdChapterInformation]]:
    """
    Extract all EBD keys from the given file.
    Returns a dictionary with all EBD keys as keys and the respective EBD titles as values.
    E.g. key: "E_0003", value: "Bestellung der Aggregationsebene RZ prüfen"
    """
    document = get_document(docx_file_path)
    result: Dict[str, Tuple[str, EbdChapterInformation]] = {}
    for paragraph, ebd_kapitel in _enrich_paragraphs_with_sections(document.paragraphs):
        match = _ebd_key_with_heading_pattern.match(paragraph.text)
        if match is None:
            continue
        ebd_key = match.groupdict()["key"]
        title = match.groupdict()["title"]
        result[ebd_key] = (title, ebd_kapitel)
        _logger.debug("Found EBD %s: '%s' (%s)", ebd_key, title, ebd_kapitel)
    _logger.info("%i EBD keys have been found", len(result))
    return result
